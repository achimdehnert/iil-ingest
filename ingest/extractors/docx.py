"""DOCX extractor using python-docx."""
from __future__ import annotations

from ingest.types import ExtractedContent

SUPPORTED_MIMES: frozenset[str] = frozenset({
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
})


class DOCXExtractor:
    """Extract text and tables from DOCX bytes (requires python-docx)."""

    supported_mimes = SUPPORTED_MIMES

    def can_handle(self, mime_type: str) -> bool:
        return mime_type in self.supported_mimes

    def extract(self, data: bytes) -> ExtractedContent:
        try:
            import docx  # type: ignore[import]
        except ImportError as exc:
            raise ImportError(
                "python-docx required. Install with: pip install iil-ingest[docx]"
            ) from exc

        import io

        errors: list[str] = []
        text_parts: list[str] = []
        tables: list[list[list[str]]] = []

        try:
            doc = docx.Document(io.BytesIO(data))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                tables.append([[cell.text for cell in row.cells] for row in table.rows])
        except Exception as exc:  # noqa: BLE001
            errors.append(f"DOCX read error: {exc}")

        return ExtractedContent(
            raw_bytes=data,
            text="\n".join(text_parts),
            tables=tables,
            metadata={},
            mime_type=next(iter(SUPPORTED_MIMES)),
            page_count=0,
            extraction_errors=errors,
        )
